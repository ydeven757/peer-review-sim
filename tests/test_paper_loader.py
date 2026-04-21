"""Tests for paper_loader module."""
import pytest
import tempfile
import os


def test_load_raw_text():
    from app.paper_loader import load_paper
    text = load_paper("This is my paper content", "text")
    assert "This is my paper content" in text


def test_load_text_with_auto_detection():
    from app.paper_loader import load_paper
    text = load_paper("This is raw content")
    assert "This is raw content" in text


def test_load_unsupported_format():
    from app.paper_loader import load_paper
    with pytest.raises(ValueError, match="Unsupported format"):
        load_paper("/tmp/paper.xyz", "xyz")


def test_load_docx(tmp_path):
    # Create a minimal .docx file
    from docx import Document

    doc = Document()
    doc.add_paragraph("Introduction")
    doc.add_paragraph("This is a test paper for peer review.")
    doc.add_paragraph("Section 1: Methodology")

    path = tmp_path / "test_paper.docx"
    doc.save(str(path))

    from app.paper_loader import load_paper
    text = load_paper(str(path), "docx")
    assert "Introduction" in text
    assert "test paper" in text
    assert "Methodology" in text


def test_load_pdf(tmp_path):
    # Create a minimal PDF using reportlab
    from reportlab.pdfgen import canvas

    path = tmp_path / "test_paper.pdf"
    c = canvas.Canvas(str(path))
    c.drawString(72, 750, "Introduction")
    c.drawString(72, 735, "This is a test paper for peer review.")
    c.drawString(72, 720, "Section 1: Methodology")
    c.save()

    from app.paper_loader import load_paper
    text = load_paper(str(path), "pdf")
    assert "Introduction" in text
    assert "test paper" in text